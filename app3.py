import streamlit as st
import os
import tempfile
import torch

from langchain_huggingface import HuggingFaceEmbeddings, HuggingFacePipeline
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableLambda
from langchain_core.output_parsers import StrOutputParser

import fitz  # PyMuPDF
from docx import Document as DocxDocument

# ---------------- CONFIG ----------------

EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"

# ---------------- RAG CORE ----------------

def build_retriever(uploaded_files):
    all_documents = []

    with tempfile.TemporaryDirectory() as temp_dir:
        for uploaded_file in uploaded_files:
            path = os.path.join(temp_dir, uploaded_file.name)
            with open(path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            ext = uploaded_file.name.lower().split(".")[-1]

            # PDF
            if ext == "pdf":
                pdf = fitz.open(path)
                for i, page in enumerate(pdf):
                    text = page.get_text("text")
                    if text.strip():
                        all_documents.append(
                            Document(
                                page_content=text,
                                metadata={
                                    "source": uploaded_file.name,
                                    "page": i + 1
                                }
                            )
                        )
                pdf.close()

            # DOCX
            elif ext == "docx":
                doc = DocxDocument(path)
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                if text:
                    all_documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source": uploaded_file.name,
                                "page": "N/A"
                            }
                        )
                    )

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )

    chunks = splitter.split_documents(all_documents)

    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL_NAME
    )

    # ✅ IN-MEMORY, SESSION-LOCAL CHROMA
    vector_db = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings
    )

    return vector_db.as_retriever(search_kwargs={"k": 8})


def build_rag_chain(retriever):
    llm = HuggingFacePipeline.from_model_id(
        model_id="google/flan-t5-small",
        task="text2text-generation",
        device=0 if torch.cuda.is_available() else -1,
        pipeline_kwargs={
            "max_new_tokens": 300,
            "temperature": 0.0
        }
    )

    prompt = ChatPromptTemplate.from_template(
        """
        You are a helpful, professional, and respectful assistant.
        Answer the question based ONLY on the provided context.
        If the answer is not in the context, say so clearly.
        Create a visually appealing, well-structured summarized answer using proper markdown formatting:

        ## Structure Your Answer:
        - **Use headings and subheadings** for organization
        - **Bold key terms** with **double asterisks**
        - **Create bullet points or numbered lists** for step-by-step information
        - **Use code blocks** with ```language for code snippets
        - **Include tables** with | separators for comparisons
        - **Add blockquotes** with > for important quotes from context
        - **Use LaTeX** for mathematical formulas: $equation$
        - **Emphasize important points** with *italics* or **bold**

        Compare the documents if multiple are relevant.
        Answer with:
        - A clear comparison between entities
        - Structured, easy-to-read format (headings, bullet points)
        - Mention if information is not present.

        Context:
        {context}

        Question:
        {question}

        Answer:
        """
    )

    def retrieve(question):
        docs = retriever.invoke(question)
        context = "\n\n".join(d.page_content for d in docs)
        return {"context": context, "question": question, "docs": docs}

    return (
        RunnableLambda(retrieve)
        | {
            "answer": prompt | llm | StrOutputParser(),
            "docs": RunnableLambda(lambda x: x["docs"])
        }
    )

# ---------------- STREAMLIT APP ----------------

def main():
    st.set_page_config(
        page_title="RAG Q&A",
        page_icon="📚",
        layout="wide"
    )

    st.title("📚 RAG Document Q&A")

    # 🔐 SESSION ISOLATION
    if "retriever" not in st.session_state:
        st.session_state.retriever = None

    if "rag_chain" not in st.session_state:
        st.session_state.rag_chain = None

    if "messages" not in st.session_state:
        st.session_state.messages = []

    with st.sidebar:
        st.header("📤 Upload Documents")

        uploaded_files = st.file_uploader(
            "Upload PDF or DOCX",
            type=["pdf", "docx"],
            accept_multiple_files=True
        )

        if st.button("🔄 Process Documents", type="primary"):
            if uploaded_files:
                st.session_state.retriever = build_retriever(uploaded_files)
                st.session_state.rag_chain = build_rag_chain(
                    st.session_state.retriever
                )
                st.session_state.messages = []
                st.success("Documents processed successfully!")
                st.rerun()
            else:
                st.error("Please upload at least one document.")

    # Chat history
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    if st.session_state.rag_chain:
        if query := st.chat_input("Ask a question about your documents"):
            st.session_state.messages.append(
                {"role": "user", "content": query}
            )

            with st.chat_message("user"):
                st.markdown(query)

            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    result = st.session_state.rag_chain.invoke(query)

                st.markdown(result["answer"])

                with st.expander("📄 Sources"):
                    for i, doc in enumerate(result["docs"], 1):
                        st.markdown(
                            f"**{i}. {doc.metadata.get('source')} "
                            f"(Page: {doc.metadata.get('page')})**"
                        )

            st.session_state.messages.append(
                {"role": "assistant", "content": result["answer"]}
            )
    else:
        st.info("👆 Upload documents and click *Process Documents* to begin.")


if __name__ == "__main__":
    main()
