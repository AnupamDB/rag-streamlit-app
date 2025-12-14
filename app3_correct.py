import streamlit as st
import os
import tempfile
import torch
import shutil

from langchain_huggingface import HuggingFaceEmbeddings, HuggingFacePipeline
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from huggingface_hub import login
from dotenv import load_dotenv
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import (
    RunnablePassthrough,
    RunnableParallel,
    RunnableLambda
)
from langchain_core.output_parsers import StrOutputParser
import fitz  # PyMuPDF
from langchain_core.documents import Document
from docx import Document as DocxDocument

load_dotenv()

# --- Configuration ---

HF_API_TOKEN = os.getenv("HUGGINGFACEHUB_API_TOKEN", "YOUR_HF_TOKEN")

LLM_MODEL_ID = "gpt2"
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"
CHROMA_PERSIST_DIR = "./chroma_db_storage"

if HF_API_TOKEN == "YOUR_HF_TOKEN":
    st.error(
        "Please set your HUGGINGFACEHUB_API_TOKEN environment variable "
        "or replace 'YOUR_HF_TOKEN' in the script."
    )
    st.stop()

# --- Core RAG Functions ---

@st.cache_resource(show_spinner=False)
def get_vector_db_and_retriever(_uploaded_files):
    """Loads, chunks, and indexes the documents into ChromaDB."""
    if not _uploaded_files:
        return None

    st.write("Processing documents...")
    all_documents = []

    with tempfile.TemporaryDirectory() as temp_dir:
        for uploaded_file in _uploaded_files:
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            try:
                file_ext = uploaded_file.name.lower().split(".")[-1]

                # ---------- PDF ----------
                if file_ext == "pdf":
                    pdf = fitz.open(temp_file_path)

                    for page_index, page in enumerate(pdf):
                        text = page.get_text("text")

                        if text.strip():
                            all_documents.append(
                                Document(
                                    page_content=text,
                                    metadata={
                                        "source": uploaded_file.name,
                                        "page": page_index + 1,
                                        "type": "pdf"
                                    }
                                )
                            )

                    pdf.close()

                # ---------- WORD (.docx) ----------
                elif file_ext == "docx":
                    doc = DocxDocument(temp_file_path)

                    full_text = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)

                    if full_text:
                        all_documents.append(
                            Document(
                                page_content="\n".join(full_text),
                                metadata={
                                    "source": uploaded_file.name,
                                    "page": "N/A",
                                    "type": "docx"
                                }
                            )
                        )

                else:
                    st.warning(f"Unsupported file type: {uploaded_file.name}")

                print(all_documents)

            except Exception as e:
                st.error(f"❌ Could not load {uploaded_file.name}: {e}")



    if not all_documents:
        return None

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    final_chunks = text_splitter.split_documents(all_documents)

    if os.path.exists(CHROMA_PERSIST_DIR):
        try:
            import time
            import gc

            gc.collect()
            time.sleep(1)
            shutil.rmtree(CHROMA_PERSIST_DIR)

        except (PermissionError, OSError):
            try:
                import subprocess
                subprocess.run(
                    ["cmd", "/c", "rmdir", "/s", "/q", CHROMA_PERSIST_DIR],
                    check=False,
                    timeout=5
                )
            except Exception:
                try:
                    for root, dirs, files in os.walk(
                        CHROMA_PERSIST_DIR, topdown=False
                    ):
                        for file in files:
                            try:
                                os.remove(os.path.join(root, file))
                            except:
                                pass
                except Exception:
                    st.warning(
                        "⚠️ Could not completely clear old database files."
                    )

    embedding_model = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL_NAME
    )

    vector_db = Chroma.from_documents(
        documents=final_chunks,
        embedding=embedding_model,
        persist_directory=CHROMA_PERSIST_DIR
    )

    retriever = vector_db.as_retriever(search_kwargs={"k": 12})

    st.success(
        f"Successfully processed {len(all_documents)} pages into "
        f"{len(final_chunks)} searchable chunks."
    )

    return retriever


@st.cache_resource(show_spinner=False)
def get_rag_chain(_retriever, _version=23):
    llm = HuggingFacePipeline.from_model_id(
        model_id="google/flan-t5-small",
        task="text2text-generation",
        device=0 if torch.cuda.is_available() else -1,
        pipeline_kwargs={
            "max_new_tokens": 300,
            "temperature": 0.0
        }
    )

    prompt = ChatPromptTemplate.from_template(
        """
        You are a helpful, professional, and respectful assistant.
        Answer the question based ONLY on the provided context.
        If the answer is not in the context, say so clearly.
        Create a visually appealing, well-structured answer using proper markdown formatting:

        ## Structure Your Answer:
        - **Use headings and subheadings** for organization
        - **Bold key terms** with **double asterisks**
        - **Create bullet points or numbered lists** for step-by-step information
        - **Use code blocks** with ```language for code snippets
        - **Include tables** with | separators for comparisons
        - **Add blockquotes** with > for important quotes from context
        - **Use LaTeX** for mathematical formulas: $equation$
        - **Emphasize important points** with *italics* or **bold**

        Compare the documents if multiple are relevant.
        Answer with:
        - A clear comparison between entities
        - Structured, easy-to-read format (headings, bullet points)
        - Mention if information is not present

        Context:
        {context}

        Question:
        {question}

        Answer:
        """
    )

    def retrieve_and_format(question):
        docs = _retriever.invoke(question)
        context = "\n\n".join(d.page_content for d in docs)
        return {
            "context": context,
            "question": question,
            "docs": docs
        }

    rag_chain = (
        RunnableLambda(retrieve_and_format)
        | {
            "answer": prompt | llm | StrOutputParser(),
            "context": RunnableLambda(lambda x: x["docs"])
        }
    )

    return rag_chain


# --- Streamlit UI ---

def main():
     # 🔥 HARD RESET ON BROWSER REFRESH
    if "session_initialized" not in st.session_state:
        # First run after browser refresh
        st.session_state.session_initialized = True

        # Clear cached resources
        get_vector_db_and_retriever.clear()
        get_rag_chain.clear()

        # Clear session state except this flag
        for key in list(st.session_state.keys()):
            if key != "session_initialized":
                del st.session_state[key]

        # Optional: delete chroma db
        if os.path.exists(CHROMA_PERSIST_DIR):
            shutil.rmtree(CHROMA_PERSIST_DIR, ignore_errors=True)

        st.rerun()

    # 🔑 STEP 2 GOES HERE (RIGHT AFTER RESET)
    if "uploader_key" not in st.session_state:
        st.session_state.uploader_key = 0


    st.set_page_config(
        page_title="RAG Document Q&A",
        page_icon="📚",
        layout="wide"
    )

    st.title("📚 RAG Document Q&A Assistant")

    with st.sidebar:
        st.header("📤 Upload Documents")
        st.markdown("Upload your PDF files to build the knowledge base.")

        uploaded_files = st.file_uploader(
            "Choose PDF or Word files",
            type=["pdf", "docx"],
            accept_multiple_files=True,
            key=f"uploader_{st.session_state.uploader_key}"
        )


        if st.button("🔄 Process Documents", type="primary"):
            get_vector_db_and_retriever.clear()
            st.session_state.messages = []

            if uploaded_files:
                get_vector_db_and_retriever(uploaded_files)
            else:
                st.error("❌ Please select at least one PDF file.")
            st.rerun()

    st.markdown(
        """
        Welcome to the **RAG Document Q&A Assistant**!
        Upload PDFs and ask questions based on them.
        """
    )

    retriever = get_vector_db_and_retriever(uploaded_files)

    if retriever:
        rag_chain = get_rag_chain(retriever, _version=22)

        if "messages" not in st.session_state:
            st.session_state.messages = []

        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        if user_query := st.chat_input("Ask a question about your documents..."):

            # User message
            st.session_state.messages.append(
                {"role": "user", "content": user_query}
            )

            with st.chat_message("user"):
                st.markdown(user_query)

            # ---- ASSISTANT (FIXED) ----
            assistant_container = st.empty()

            with assistant_container.container():
                with st.chat_message("assistant"):
                    with st.spinner("🔍 Searching and generating answer..."):
                        response = rag_chain.invoke(user_query)

                    answer = response["answer"]
                    sources = response["context"]

                    st.markdown(answer)

                    with st.expander("📄 Sources Referenced"):
                        if sources:
                            for i, doc in enumerate(sources, 1):
                                source_name = (
                                    doc.metadata.get("source", "Unknown File")
                                    .split("/")[-1]
                                )
                                page_num = doc.metadata.get("page", "N/A")

                                st.markdown(
                                    f"**{i}. {source_name}** (Page: {page_num})"
                                )

                                snippet = (
                                    doc.page_content[:200] + "..."
                                    if len(doc.page_content) > 200
                                    else doc.page_content
                                )
                                st.caption(snippet)
                        else:
                            st.write("No sources found.")

            # Save AFTER render
            st.session_state.messages.append(
                {"role": "assistant", "content": answer}
            )



    else:
        st.info(
            "👆 Please upload PDF documents and process them to start."
        )


if __name__ == "__main__":
    main()
